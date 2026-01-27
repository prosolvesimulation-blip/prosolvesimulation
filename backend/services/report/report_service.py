import os
import sys
import json
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# ==============================================================================
# REPORT_SERVICE.PY - TECHNICAL DOCUMENTATION GENERATOR
# Goal: Create professional customizable Word reports from simulation data.
# Stack: python-docx
# ==============================================================================

class TechnicalReportGenerator:
    def __init__(self, project_path):
        self.project_path = project_path
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Configures fonts and professional styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Heading 1
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(16)
        h1.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue
        h1.font.bold = True
        
        # Heading 2
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(14)
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.font.bold = True

    def generate_report(self, data_context, selection=None):
        """
        Main orchestration method.
        selection: dict of bools indicating which sections to include.
        data_context: {
            "project_name": str,
            "project_config": dict (from project.json),
            "mass_props": dict,
            "reactions": list,
            "max_stress": float,
            "images": [path1, path2]
        }
        """
        if selection is None:
            selection = {k: True for k in ["model", "materials", "geometries", "restrictions", "loads", "load_cases", "results"]}

        self._add_title_page(data_context)
        self.doc.add_page_break()
        
        self._add_executive_summary(data_context)

        # 1. Model Properties (Mass/CG)
        if selection.get("model", True):
            self._add_model_properties(data_context)

        # 2. Project Configuration
        config = data_context.get("project_config", {})
        
        if selection.get("materials", True) and config.get("materials"):
            self._add_materials_info(config["materials"])
            
        if selection.get("geometries", True) and config.get("geometries"):
            self._add_geometries_info(config["geometries"])
            
        if selection.get("restrictions", True) and config.get("restrictions"):
            self._add_restrictions_info(config["restrictions"])
            
        if selection.get("loads", True) and config.get("loads"):
            self._add_loads_info(config["loads"])
            
        if selection.get("load_cases", True) and config.get("load_cases"):
            self._add_load_cases_info(config["load_cases"])

        # 3. Structural Results
        if selection.get("results", True):
            self._add_structural_results(data_context)
            if data_context.get("images"):
                self._add_visual_appendix(data_context["images"])
            
        return self.doc

    def _add_materials_info(self, materials):
        self.doc.add_heading('2.2 Materials', level=2)
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'E (MPa)'
        hdr_cells[2].text = 'Nu'
        hdr_cells[3].text = 'Density (kg/m3)'
        
        for mat in materials:
            row = table.add_row().cells
            row[0].text = str(mat.get('name', ''))
            row[1].text = str(mat.get('young', ''))
            row[2].text = str(mat.get('poisson', ''))
            row[3].text = str(mat.get('rho', ''))

    def _add_geometries_info(self, geometries):
        self.doc.add_heading('2.3 Geometries', level=2)
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Group'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Section/Details'
        
        for geom in geometries:
            row = table.add_row().cells
            row[0].text = str(geom.get('group', ''))
            row[1].text = str(geom.get('_category', ''))
            
            details = ""
            if geom.get('section_type'):
                details = f"Section: {geom.get('section_type')}"
            elif geom.get('section_params', {}).get('thickness'):
                details = f"Thickness: {geom.get('section_params', {}).get('thickness')} mm"
            row[2].text = details

    def _add_restrictions_info(self, restrictions):
        self.doc.add_heading('2.4 Restrictions', level=2)
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Group'
        hdr_cells[2].text = 'Blocked DDLs'
        
        for res in restrictions:
            row = table.add_row().cells
            row[0].text = str(res.get('name', ''))
            row[1].text = str(res.get('group', ''))
            
            ddls = []
            if res.get('dx'): ddls.append('DX')
            if res.get('dy'): ddls.append('DY')
            if res.get('dz'): ddls.append('DZ')
            if res.get('drx'): ddls.append('DRX')
            if res.get('dry'): ddls.append('DRY')
            if res.get('drz'): ddls.append('DRZ')
            row[2].text = ", ".join(ddls)

    def _add_loads_info(self, loads):
        self.doc.add_heading('2.5 Applied Loads', level=2)
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Group/Type'
        hdr_cells[2].text = 'Parameters'
        hdr_cells[3].text = 'Value'
        
        for load in loads:
            row = table.add_row().cells
            row[0].text = str(load.get('name', ''))
            row[1].text = f"{load.get('group', 'Global')} ({load.get('type', '')})"
            
            params = []
            val = ""
            if load.get('type') == 'PESANTEUR':
                params = ['Direction']
                dirs = []
                if load.get('dx'): dirs.append(f"X: {load.get('dx')}")
                if load.get('dy'): dirs.append(f"Y: {load.get('dy')}")
                if load.get('dz'): dirs.append(f"Z: {load.get('dz')}")
                val = ", ".join(dirs)
            else:
                if load.get('fx'): params.append(f"FX: {load.get('fx')} N")
                if load.get('fy'): params.append(f"FY: {load.get('fy')} N")
                if load.get('fz'): params.append(f"FZ: {load.get('fz')} N")
                val = "Point Force" if load.get('type') == 'FORCE_NODA' else "Pressure"
            
            row[2].text = ", ".join(params)
            row[3].text = val

    def _add_load_cases_info(self, load_cases):
        self.doc.add_heading('2.6 Load Cases', level=2)
        for lc in load_cases:
            p = self.doc.add_paragraph()
            p.add_run(f"Case: {lc.get('name', 'Unnamed')}").bold = True
            p.add_run(f"\nIncluded Loads: {', '.join(lc.get('loads', []))}")
            p.add_run(f"\nIncluded Restrictions: {', '.join(lc.get('restrictions', []))}")

    def _add_title_page(self, data):
        """Creates a professional cover page."""
        # Logo Placeholder (Text for now)
        title = self.doc.add_heading('TECHNICAL SIMULATION REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\nProject: {data.get('project_name', 'Unnamed Project')}\n")
        run.font.size = Pt(24)
        run.bold = True
        
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(f"Generated by ProSolve Professional\n{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    def _add_executive_summary(self, data):
        self.doc.add_heading('1. Executive Summary', level=1)
        self.doc.add_paragraph(
            f"This report summarizes the structural analysis performed on {data.get('project_name')}. "
            f"The simulation evaluated mechanical integrity using Finite Element Analysis (FEA) via Code_Aster."
        )
        
    def _add_model_properties(self, data):
        self.doc.add_heading('2. Model Properties', level=1)
        
        mass_data = data.get("mass_props", {})
        if mass_data:
            self.doc.add_heading('2.1 Mass & Inertia', level=2)
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Property'
            hdr_cells[1].text = 'Value'
            
            props = [
                ("Total Mass (kg)", f"{mass_data.get('mass', 0):.4f}"),
                ("Center of Gravity X (mm)", f"{mass_data.get('cdg_x', 0):.2f}"),
                ("Center of Gravity Y (mm)", f"{mass_data.get('cdg_y', 0):.2f}"),
                ("Center of Gravity Z (mm)", f"{mass_data.get('cdg_z', 0):.2f}")
            ]
            
            for k, v in props:
                row_cells = table.add_row().cells
                row_cells[0].text = k
                row_cells[1].text = v

    def _add_structural_results(self, data):
        self.doc.add_heading('3. Structural Results', level=1)
        
        # Max Stress
        self.doc.add_heading('3.1 Von Mises Stress', level=2)
        p = self.doc.add_paragraph()
        p.add_run("Maximum Identified Stress: ").bold = True
        p.add_run(f"{data.get('max_stress', 0):.2f} MPa")
        
        # Reactions
        reactions = data.get("reactions", [])
        if reactions:
            self.doc.add_heading('3.2 Reaction Forces', level=2)
            for r in reactions:
                self.doc.add_paragraph(f"Case: {r.get('case_name', 'Unknown')}", style='Body Text')
                table = self.doc.add_table(rows=2, cols=6)
                table.style = 'Light List'
                
                # Header
                hdr = table.rows[0].cells
                hdr[0].text = 'Fx (N)'
                hdr[1].text = 'Fy (N)'
                hdr[2].text = 'Fz (N)'
                hdr[3].text = 'Mx (N.m)'
                hdr[4].text = 'My (N.m)'
                hdr[5].text = 'Mz (N.m)'
                
                # Values
                row = table.rows[1].cells
                row[0].text = f"{r.get('fx', 0):.2f}"
                row[1].text = f"{r.get('fy', 0):.2f}"
                row[2].text = f"{r.get('fz', 0):.2f}"
                row[3].text = f"{r.get('mx', 0):.2f}"
                row[4].text = f"{r.get('my', 0):.2f}"
                row[5].text = f"{r.get('mz', 0):.2f}"
                self.doc.add_paragraph("") # Spacing

    def _add_visual_appendix(self, images):
        """Adds user visualisations to the end."""
        self.doc.add_page_break()
        self.doc.add_heading('Appendix: Analysis Visuals', level=1)
        
        for img_path in images:
            if os.path.exists(img_path):
                try:
                    self.doc.add_picture(img_path, width=Inches(6.0))
                    self.doc.add_paragraph("Figure: User Visualization", style='Caption')
                except Exception as e:
                    self.doc.add_paragraph(f"[Image Error: {os.path.basename(img_path)}]")

def create_report_file(project_path, context, selection=None, output_name="Structural_Report.docx"):
    """Entry point for API."""
    try:
        report = TechnicalReportGenerator(project_path)
        doc = report.generate_report(context, selection=selection)
        
        out_path = os.path.join(project_path, output_name)
        doc.save(out_path)
        return {"status": "success", "file_path": out_path}
    except Exception as e:
        return {"status": "error", "message": str(e)}
